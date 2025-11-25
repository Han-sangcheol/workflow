"""
DOCX 파일 파싱 모듈
python-docx를 사용하여 DOCX 파일의 텍스트를 추출합니다.
"""

import logging
from pathlib import Path
from typing import Optional

try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)


class DOCXParser:
    """DOCX 문서 텍스트 추출 클래스"""

    @staticmethod
    def is_available() -> bool:
        """python-docx 설치 여부 확인"""
        return Document is not None

    def parse_file(self, file_path: str) -> Optional[str]:
        """
        DOCX 파일에서 텍스트 추출
        
        Args:
            file_path: DOCX 파일 경로
            
        Returns:
            추출된 텍스트 (실패 시 None)
        """
        if not self.is_available():
            logger.error("python-docx가 설치되지 않았습니다")
            return None

        try:
            path = Path(file_path)
            if not path.exists():
                logger.error(f"파일을 찾을 수 없습니다: {file_path}")
                return None

            logger.info(f"DOCX 파일 파싱 시작: {path.name}")
            
            doc = Document(file_path)
            text_parts = []
            
            # 단락 추출
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # 표 추출
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_parts.append(f"\n[표]\n{table_text}")
            
            result = "\n".join(text_parts)
            logger.info(
                f"DOCX 파싱 완료: {len(doc.paragraphs)} 단락, "
                f"{len(doc.tables)} 표, {len(result)} 문자"
            )
            
            return result if result.strip() else None

        except Exception as e:
            logger.error(f"DOCX 파싱 오류: {str(e)}")
            return None

    def _extract_table_text(self, table) -> str:
        """표에서 텍스트 추출"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)

