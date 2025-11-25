"""
출력 생성 유틸리티
분석 결과를 DOCX 파일로 저장합니다.
"""

import logging
from pathlib import Path
from datetime import datetime
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

logger = logging.getLogger(__name__)


class OutputGenerator:
    """결과 저장 클래스"""

    @staticmethod
    def is_available() -> bool:
        """python-docx 설치 여부 확인"""
        return Document is not None

    def save_summary(
        self,
        content: str,
        output_path: str
    ) -> bool:
        """
        통합 회의록 저장
        
        Args:
            content: 회의록 내용
            output_path: 저장 경로
            
        Returns:
            성공 여부
        """
        return self._save_document(
            title="팀 업무 통합 회의록",
            content=content,
            output_path=output_path
        )

    def save_thanks(
        self,
        content: str,
        output_path: str
    ) -> bool:
        """
        감사 인사 저장
        
        Args:
            content: 감사 인사 내용
            output_path: 저장 경로
            
        Returns:
            성공 여부
        """
        return self._save_document(
            title="팀원 감사 인사",
            content=content,
            output_path=output_path
        )

    def _save_document(
        self,
        title: str,
        content: str,
        output_path: str
    ) -> bool:
        """
        DOCX 문서 저장
        
        Args:
            title: 문서 제목
            content: 문서 내용
            output_path: 저장 경로
            
        Returns:
            성공 여부
        """
        if not self.is_available():
            logger.error("python-docx가 설치되지 않았습니다")
            return False

        try:
            doc = Document()
            
            # 제목 추가
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 날짜 추가
            date_para = doc.add_paragraph(
                datetime.now().strftime("%Y년 %m월 %d일")
            )
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 구분선
            doc.add_paragraph("=" * 50)
            
            # 내용 추가
            for line in content.split("\n"):
                if line.strip():
                    para = doc.add_paragraph(line)
                    # 폰트 설정
                    for run in para.runs:
                        run.font.size = Pt(11)
                else:
                    doc.add_paragraph()
            
            # 파일 저장
            doc.save(output_path)
            logger.info(f"문서 저장 완료: {output_path}")
            return True

        except Exception as e:
            logger.error(f"문서 저장 오류: {str(e)}")
            return False

    @staticmethod
    def generate_default_filename(prefix: str) -> str:
        """
        기본 파일명 생성
        
        Args:
            prefix: 파일명 접두사
            
        Returns:
            생성된 파일명
        """
        date_str = datetime.now().strftime("%y%m%d")
        return f"{prefix}_{date_str}.docx"

